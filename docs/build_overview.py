from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(4)

NAVY = RGBColor(0x01, 0x24, 0x40)


def heading(text, size=14, space_before=6, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = NAVY
    return p


def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    return p


def para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    return p


# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("TTfE Website — Briefing Note")
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = NAVY
title.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sr = sub.add_run("ttfe.org.uk  |  internal briefing for the TTfE team")
sr.italic = True
sr.font.size = Pt(9)
sub.paragraph_format.space_after = Pt(8)

# Purpose
subheading("Purpose")
para(
    "This note summarises how the Tram Trains for Edinburgh website is built, hosted and "
    "maintained, the reasons behind the technical choices, and the ongoing costs. It is "
    "aimed at committee members and volunteers who are not closely involved in the "
    "technical side of the campaign."
)

# How it works
subheading("How the website works")
para(
    "The website is a set of plain web pages (HTML). There is no database or server "
    "program to keep running — just files that are delivered to visitors. Content is "
    "written in simple text files and automatically turned into styled pages by a tool "
    "called Eleventy each time a change is made."
)
para(
    "News posts are managed through Decap CMS, a visual editor reached at "
    "ttfe.org.uk/admin/. Editors log in with their GitHub account; posts save back to "
    "our code repository and the site rebuilds itself within a minute or so."
)

# Hosting
subheading("How it is hosted")
para(
    "The site is hosted on Cloudflare Pages. Our source code lives on GitHub under the "
    "Tram-Trains-for-Edinburgh organisation. Every change pushed to GitHub (by an "
    "editor, a volunteer, or the CMS) triggers Cloudflare to rebuild and publish the "
    "site automatically, worldwide, within roughly a minute."
)
para(
    "Form submissions (contact and membership) are handled by small serverless "
    "functions running on Cloudflare, which use the Resend email service to deliver "
    "each submission to the campaign inbox. Cloudflare Turnstile is used to block spam."
)

# Tech stack rationale
subheading("Why this technology was chosen")
bullet(
    "Longevity: static sites can be handed to a new maintainer with minimal fuss and "
    "will keep working for years without attention."
)
bullet(
    "No lock-in to any one supplier: the site can be moved to a different host in a "
    "matter of hours if Cloudflare ever becomes unsuitable."
)
bullet(
    "Campaign-friendly editing: non-technical volunteers can add news posts through a "
    "browser without touching code."
)
bullet(
    "Security: with no database and no server program, there is very little to attack "
    "or patch."
)
bullet(
    "Predictable cost: static hosting at our scale is free. We do not risk a surprise "
    "bill driven by traffic or usage."
)

# Non-technical strengths
subheading("Strengths in plain terms")
bullet(
    "Fast and reliable: served from Cloudflare's global network, pages load quickly "
    "and the site stays up even under sudden traffic (e.g. press coverage)."
)
bullet(
    "Always up to date: edits appear on the live site within about a minute."
)
bullet(
    "Accountable: every change is logged on GitHub with the editor's name and time, "
    "providing an audit trail."
)
bullet(
    "Portable: all content, styling and code are in files we own. Nothing is trapped "
    "inside a proprietary platform."
)

# Openness
subheading("Open and transparent")
para(
    "The source code and content for the website are maintained in a public GitHub "
    "repository at https://github.com/Tram-Trains-for-Edinburgh/tram-trains. Anyone "
    "can inspect how the site is built, review its history, or raise issues. "
    "Publishing the repository openly reinforces the campaign's transparent, "
    "community-led, community-driven approach: there are no hidden systems, and the "
    "site could be picked up and continued by any supporter with the relevant skills."
)

# Domain choice
subheading("Domain choice (ttfe.org.uk, not ttfe.co.uk)")
para(
    "The campaign uses the .org.uk domain because TTfE is a volunteer, non-commercial "
    "organisation. In the UK, .co.uk is associated with commercial enterprises, whereas "
    ".org.uk is the conventional suffix for campaign groups, charities, and community "
    "organisations. The .org.uk domain better signals the campaign's non-profit "
    "character and helps build trust with the public, politicians, and press."
)
para(
    "It also signals intent to donors and to platforms that list or support non-profits "
    "(for example Goodstack, and the Charities Aid Foundation in Scotland), which "
    "generally expect non-commercial branding and a non-commercial domain."
)
para(
    "To avoid losing anyone who still has the old address bookmarked or printed on "
    "earlier material, all traffic to ttfe.co.uk will be redirected automatically to "
    "ttfe.org.uk. Visitors typing, clicking, or searching for the .co.uk domain will "
    "arrive at the correct site without needing to do anything."
)

# Costs
subheading("Costs")
para(
    "The website currently operates at no cost. All services used are on free tiers "
    "that comfortably exceed our anticipated usage:"
)
bullet("Cloudflare Pages (hosting): free — unlimited bandwidth at our traffic levels.")
bullet("Cloudflare Workers (CMS login): free tier covers 100,000 requests per day.")
bullet("GitHub (code storage): free for public repositories.")
bullet("Resend (transactional email for forms): free tier of 3,000 emails per month.")
bullet("Domain name (ttfe.org.uk): approximately £10 per year.")
para(
    "If volumes ever grew significantly (e.g. thousands of form submissions per month) "
    "some paid tiers might be needed, but these start at only a few pounds per month."
)

# Risks / dependencies
subheading("Dependencies and risks")
bullet(
    "Cloudflare account: the site, its DNS, and the CMS login all run through one "
    "Cloudflare account. Access to that account should be held by at least two "
    "trusted people."
)
bullet(
    "GitHub organisation: the code repository is owned by the "
    "Tram-Trains-for-Edinburgh GitHub organisation; at least two committee members "
    "should hold Owner roles on the organisation."
)
bullet(
    "Domain registration: ttfe.org.uk must be renewed annually; a lapse would take "
    "the site offline."
)
bullet(
    "No loss of editorial control: because the code is on GitHub, the site could be "
    "rebuilt on a different host by any competent developer."
)

# Who does what
subheading("Who does what")
bullet("Editorial (news posts): any committee member granted CMS access via GitHub.")
bullet("Page content edits: handled by the technical volunteer(s) through GitHub.")
bullet(
    "Hosting, DNS, and domain: managed centrally via the TTfE Cloudflare account."
)
bullet("Form submissions land in the TTfE campaign inbox.")

# Footer
footer = doc.add_paragraph()
fr = footer.add_run("Prepared for the TTfE committee. Please treat as an internal working document.")
fr.italic = True
fr.font.size = Pt(8)
footer.paragraph_format.space_before = Pt(10)

doc.save(r"f:/git/tram-trains/docs/ttfe-website-overview.docx")
print("Saved.")
