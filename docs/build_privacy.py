"""Build privacy-policy.docx from pf-site/src/pages/privacy.md verbatim."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("f:/git/tram-trains/pf-site/src/pages/privacy.md")
OUT = Path("f:/git/tram-trains/docs/ttfe-privacy-policy.docx")

NAVY = RGBColor(0x01, 0x24, 0x40)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)


def add_runs_with_bold(paragraph, text):
    """Render inline **bold** markup inside a paragraph, literal otherwise."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def heading(text, level):
    p = doc.add_paragraph()
    sizes = {1: 20, 2: 16, 3: 13, 4: 11.5}
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes[level])
    run.font.color.rgb = NAVY


def paragraph(text):
    p = doc.add_paragraph()
    add_runs_with_bold(p, text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    add_runs_with_bold(p, text)
    p.paragraph_format.space_after = Pt(2)


def footer(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


lines = SRC.read_text(encoding="utf-8").splitlines()
in_frontmatter = False
i = 0
if lines and lines[0].strip() == "---":
    # skip frontmatter
    i = 1
    while i < len(lines) and lines[i].strip() != "---":
        i += 1
    i += 1  # past closing ---

while i < len(lines):
    line = lines[i].rstrip()
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    # HTML footer paragraph
    if stripped.startswith("<p"):
        # strip tags, keep text
        text = re.sub(r"<[^>]+>", "", stripped).strip()
        footer(text)
        i += 1
        continue

    # Headings
    m = re.match(r"^(#{2,4})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1)) - 1  # ## -> 1, ### -> 2, #### -> 3
        heading(m.group(2), level)
        i += 1
        continue

    # Bullets
    if stripped.startswith("- "):
        bullet(stripped[2:])
        i += 1
        continue

    # Paragraph (may wrap? markdown here is single-line paragraphs)
    paragraph(stripped)
    i += 1

doc.save(OUT)
print(f"Saved: {OUT}")
